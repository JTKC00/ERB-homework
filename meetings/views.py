import os
import csv
import re
import tempfile
import shutil
from io import StringIO
from datetime import timedelta
from django.shortcuts import render, redirect, get_object_or_404
from .models import Meeting, AgendaItem, Attachment
from docx import Document
from django.core.files.storage import default_storage
from django.http import HttpResponse, HttpResponseBadRequest
from .resources import MeetingWithDetailsResource
from tablib import Dataset
from django.core.paginator import Paginator
import zipfile
from django.core.exceptions import ObjectDoesNotExist
from django.utils.text import slugify
from django.db.models import Q
from django.contrib import messages

# Create your views here.

def meeting_detail(request, meeting_id):
    meeting = Meeting.objects.get(id=meeting_id)
    agenda_items = AgendaItem.objects.filter(meeting=meeting)
    attachments = Attachment.objects.filter(meeting=meeting)
    return render(request, 'meetings/meeting_detail.html', {
        'meeting': meeting,
        'agenda_items': agenda_items,
        'attachments': attachments,
    })

def upload_attachment(request, meeting_id):
    if request.method == 'POST' and request.FILES.get('attachment_file'):
        meeting = Meeting.objects.get(id=meeting_id)
        uploaded_file = request.FILES['attachment_file']
        
        # 創建有意義的檔案名稱：日期_會議標題_原檔名
        date_str = meeting.date.strftime('%Y%m%d')
        title_slug = slugify(meeting.title)[:30]  # 限制標題長度
        original_name = uploaded_file.name
        file_extension = os.path.splitext(original_name)[1]
        meaningful_filename = f"{date_str}_{title_slug}_{original_name}"
        
        attachment = Attachment(meeting=meeting)
        attachment.file.save(meaningful_filename, uploaded_file, save=True)
        return redirect('meeting_detail', meeting_id=meeting.id)
    return render(request, 'meetings/upload.html', {'meeting_id': meeting_id})

def export_meetings(request):
    # 創建臨時目錄
    with tempfile.TemporaryDirectory() as temp_dir:
        # 1. 生成 CSV 檔案
        meeting_resource = MeetingWithDetailsResource()
        dataset = meeting_resource.export()
        csv_file_path = os.path.join(temp_dir, 'meetings.csv')
        with open(csv_file_path, 'w', encoding='utf-8') as f:
            f.write(dataset.csv)
        
        # 2. 複製所有附件檔案到臨時目錄
        attachments_dir = os.path.join(temp_dir, 'attachments')
        os.makedirs(attachments_dir, exist_ok=True)
        
        all_attachments = Attachment.objects.all()
        for attachment in all_attachments:
            try:
                source_file = attachment.file.path
                if os.path.exists(source_file):
                    # 保持原始檔案名稱
                    dest_file = os.path.join(attachments_dir, os.path.basename(attachment.file.name))
                    shutil.copy2(source_file, dest_file)
            except Exception as e:
                continue  # 跳過無法複製的檔案
        
        # 3. 創建 ZIP 檔案
        zip_file_path = os.path.join(temp_dir, 'meetings_export.zip')
        with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # 加入 CSV 檔案
            zipf.write(csv_file_path, 'meetings.csv')
            
            # 加入所有附件檔案
            for root, dirs, files in os.walk(attachments_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.join('attachments', file)
                    zipf.write(file_path, arcname)
        
        # 4. 回傳 ZIP 檔案
        with open(zip_file_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/zip')
            response['Content-Disposition'] = 'attachment; filename="meetings_export.zip"'
            return response

def import_meetings(request):
    if request.method == 'POST' and request.FILES.get('file'):
        uploaded_file = request.FILES['file']
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        try:
            if file_extension == '.zip':
                # 處理 ZIP 檔案
                with tempfile.TemporaryDirectory() as temp_dir:
                    # 1. 儲存上傳的 ZIP 檔案
                    zip_path = os.path.join(temp_dir, 'uploaded.zip')
                    with open(zip_path, 'wb') as f:
                        for chunk in uploaded_file.chunks():
                            f.write(chunk)
                    
                    # 2. 解壓縮 ZIP 檔案
                    with zipfile.ZipFile(zip_path, 'r') as zipf:
                        zipf.extractall(temp_dir)
                    
                    # 3. 尋找 CSV 檔案
                    csv_file_path = None
                    for file_name in ['meetings.csv', 'meetings_with_details.csv']:
                        potential_path = os.path.join(temp_dir, file_name)
                        if os.path.exists(potential_path):
                            csv_file_path = potential_path
                            break
                    
                    if not csv_file_path:
                        return render(request, 'import/import_error.html', {'errors': ['ZIP 檔案中找不到 meetings.csv 檔案']})
                    
                    # 4. 處理 CSV 檔案
                    imported_count = 0
                    with open(csv_file_path, 'r', encoding='utf-8') as f:
                        csv_reader = csv.reader(f)
                        headers = next(csv_reader)  # 跳過標題行
                        
                        for row in csv_reader:
                            if len(row) >= 7:  # 確保有足夠的欄位
                                try:
                                    meeting, created = Meeting.objects.get_or_create(
                                        title=row[1][:490],  # title
                                        date=row[2],         # date
                                        defaults={
                                            'start_time': row[3],
                                            'end_time': row[4],
                                            'location': row[5][:90],
                                            'attendees': row[6],
                                            'minutes': row[7] if len(row) > 7 else ''
                                        }
                                    )
                                    
                                    if created:
                                        imported_count += 1
                                        
                                        # 5. 處理議程項目（第8欄）
                                        if len(row) > 8 and row[8].strip() and row[8].strip() != '無議程項':
                                            agenda_text = row[8].strip()
                                            AgendaItem.objects.filter(meeting=meeting).delete()
                                            
                                            agenda_items = agenda_text.split(';')
                                            for agenda_item in agenda_items:
                                                agenda_item = agenda_item.strip()
                                                if agenda_item:
                                                    match = re.match(r'(\d+)\.\s*(.+?)\s*\(負責人:\s*(.+?)\)', agenda_item)
                                                    if match:
                                                        item_number = int(match.group(1))
                                                        item_title = match.group(2).strip()[:490]
                                                        responsible_person = match.group(3).strip()[:90]
                                                        
                                                        AgendaItem.objects.create(
                                                            meeting=meeting,
                                                            item_number=item_number,
                                                            item_title=item_title,
                                                            description=item_title,
                                                            responsible_person=responsible_person,
                                                            estimated_time=timedelta(minutes=30)
                                                        )
                                        
                                        # 6. 處理附件（第9欄）
                                        if len(row) > 9 and row[9].strip() and row[9].strip() != '無附件':
                                            attachment_text = row[9].strip()
                                            Attachment.objects.filter(meeting=meeting).delete()
                                            
                                            # 解析附件資訊：格式為 "filename.txt (type); filename2.txt (type2)"
                                            attachment_items = attachment_text.split(';')
                                            for attachment_item in attachment_items:
                                                attachment_item = attachment_item.strip()
                                                if attachment_item:
                                                    # 解析檔案名稱和類型
                                                    match = re.match(r'(.+?)\s*\((.+?)\)', attachment_item)
                                                    if match:
                                                        file_name = match.group(1).strip()
                                                        file_type = match.group(2).strip()
                                                        
                                                        # 尋找實際檔案
                                                        attachment_file_path = os.path.join(temp_dir, 'attachments', file_name)
                                                        if os.path.exists(attachment_file_path):
                                                            # 複製檔案到 media 目錄
                                                            from django.core.files import File
                                                            with open(attachment_file_path, 'rb') as f:
                                                                django_file = File(f, name=file_name)
                                                                Attachment.objects.create(
                                                                    meeting=meeting,
                                                                    file=django_file,
                                                                    file_type=file_type[:40]
                                                                )
                                
                                except Exception as e:
                                    continue  # 跳過有問題的行
                    
                    return render(request, 'import/import_success.html', {'imported_count': imported_count})
            
            elif file_extension in ['.csv']:
                # 處理純 CSV 檔案（保持原有邏輯）
                content = uploaded_file.read().decode('utf-8')
                lines = content.splitlines()
                imported_count = 0
                
                for line in lines[1:]:  # 跳過標題行
                    if line.strip():
                        csv_reader = csv.reader(StringIO(line))
                        try:
                            parts = next(csv_reader)
                            if len(parts) >= 7:
                                meeting, created = Meeting.objects.get_or_create(
                                    title=parts[1].strip()[:490],
                                    date=parts[2].strip(),
                                    defaults={
                                        'start_time': parts[3].strip(),
                                        'end_time': parts[4].strip(),
                                        'location': parts[5].strip()[:90],
                                        'attendees': parts[6].strip(),
                                        'minutes': parts[7].strip() if len(parts) > 7 else ''
                                    }
                                )
                                if created:
                                    imported_count += 1
                        except Exception as e:
                            continue
                                
                return render(request, 'import/import_success.html', {'imported_count': imported_count})

            elif file_extension in ['.txt', '.docx']:
                meetings_data = []
                if file_extension == '.txt':
                    content = uploaded_file.read().decode('utf-8')
                    lines = content.splitlines()
                    for line in lines:
                        if line.strip():
                            parts = line.split(',')
                            if len(parts) >= 9:
                                meetings_data.append({
                                    'title': parts[0].strip(),
                                    'date': parts[1].strip(),
                                    'start_time': parts[2].strip(),
                                    'end_time': parts[3].strip(),
                                    'location': parts[4].strip(),
                                    'attendees': parts[5].strip(),
                                    'minutes': parts[6].strip() if len(parts) > 6 else '',
                                    'Agenda Items': parts[7].strip() if len(parts) > 7 else '',
                                    'Attachments': parts[8].strip() if len(parts) > 8 else ''
                                })
                elif file_extension == '.docx':
                    doc = Document(uploaded_file)
                    for para in doc.paragraphs:
                        if para.text.strip():
                            parts = para.text.split(',')
                            if len(parts) >= 6:
                                meetings_data.append({
                                    'title': parts[0].strip(),
                                    'date': parts[1].strip(),
                                    'start_time': parts[2].strip(),
                                    'end_time': parts[3].strip(),
                                    'location': parts[4].strip(),
                                    'attendees': parts[5].strip(),
                                    'minutes': parts[6].strip() if len(parts) > 6 else ''
                                })

                for data in meetings_data:
                    meeting, created = Meeting.objects.update_or_create(
                        title=data['title'],
                        defaults={
                            'date': data['date'],
                            'start_time': data['start_time'],
                            'end_time': data['end_time'],
                            'location': data['location'],
                            'attendees': data['attendees'],
                            'minutes': data['minutes']
                        }
                    )
                    if created and 'Agenda Items' in data and data['Agenda Items']:
                        for i, item_str in enumerate(data['Agenda Items'].split(';'), 1):
                            parts = item_str.split('(')
                            if len(parts) >= 2:
                                item_title = parts[0].strip()
                                responsible = parts[1].replace(')', '').strip()
                                AgendaItem.objects.create(
                                    meeting=meeting,
                                    item_number=i,
                                    item_title=item_title,
                                    description='',
                                    responsible_person=responsible,
                                    estimated_time=timedelta(minutes=30)
                                )
                    if created and 'Attachments' in data and data['Attachments']:
                        for att_name in data['Attachments'].split(';'):
                            Attachment.objects.create(
                                meeting=meeting,
                                file=f"attachments/{att_name.strip()}",
                                file_type=os.path.splitext(att_name.strip())[1].replace('.', '')
                            )

                return render(request, 'import/import_success.html')

            else:
                return render(request, 'import/import_error.html', {'errors': ['不支持的文件格式。']})

        except Exception as e:
            return render(request, 'import/import_error.html', {'errors': [str(e)]})

    return render(request, 'import/import_form.html')

def meeting_detail(request, meeting_id):
    meeting = Meeting.objects.get(id=meeting_id)
    agenda_items = AgendaItem.objects.filter(meeting=meeting)
    attachments = Attachment.objects.filter(meeting=meeting)
    return render(request, 'meetings/meeting_detail.html', {
        'meeting': meeting,
        'agenda_items': agenda_items,
        'attachments': attachments,
    })

def meeting_list(request):
    # 初始查詢集
    meetings = Meeting.objects.all()

    # 搜索功能
    search_query = request.GET.get('q')
    if search_query:
        meetings = meetings.filter(
            Q(title__icontains=search_query) | Q(attendees__icontains=search_query)
        )

    # 過濾功能
    date_filter = request.GET.get('date')
    if date_filter:
        meetings = meetings.filter(date=date_filter)

    location_filter = request.GET.get('location')
    if location_filter:
        meetings = meetings.filter(location__icontains=location_filter)

    # 排序功能
    sort_by = request.GET.get('sort', '-date')  # 默認按日期降序
    if sort_by in ['date', '-date', 'title', '-title']:
        meetings = meetings.order_by(sort_by)
    else:
        meetings = meetings.order_by('-date')  # 默認排序

    # 分頁功能
    paginator = Paginator(meetings, 15)  # 每頁15個會議
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # 傳遞參數到模板
    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'date_filter': date_filter,
        'location_filter': location_filter,
        'sort_by': sort_by,
    }
    return render(request, 'meetings/meeting_list.html', context)

def delete_meeting(request, meeting_id):
    meeting = get_object_or_404(Meeting, id=meeting_id)
    if request.method == 'POST':
        meeting.delete()
        messages.success(request, '會議已成功刪除。')
        return redirect('meetings:meeting_list')
    return render(request, 'meetings/confirm_delete.html', {'meeting': meeting})

def delete_selected_meetings(request):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_meetings')
        if not selected_ids:
            return HttpResponseBadRequest("請選擇至少一個會議。")
        meetings = Meeting.objects.filter(id__in=selected_ids)
        if meetings.exists():
            meetings.delete()
            messages.success(request, f'成功刪除了 {meetings.count()} 個會議。')
        return redirect('meetings:meeting_list')
    return HttpResponseBadRequest("無效請求。")

def delete_all_meetings(request):
    if request.method == 'POST':
        Meeting.objects.all().delete()  # 刪除所有會議（級聯刪除 AgendaItem 和 Attachment）
        messages.success(request, '所有會議已成功刪除。')
        return redirect('meetings:meeting_list')
    return HttpResponseBadRequest("無效請求，必須使用 POST 方法。")

def generate_meeting_doc(request, meeting_id):
    try:
        # 獲取會議數據
        meeting = Meeting.objects.get(id=meeting_id)
        agenda_items = AgendaItem.objects.filter(meeting=meeting)
        attachments = Attachment.objects.filter(meeting=meeting)

        # 創建臨時目錄
        temp_dir = default_storage.path('temp_' + str(meeting_id))
        os.makedirs(temp_dir, exist_ok=True)

        # 格式化文件名：使用日期和標題
        date_str = meeting.date.strftime('%Y-%m-%d')
        title_slug = slugify(meeting.title)[:50]  # 限制長度，避免文件名過長
        base_name = f'{date_str}_{title_slug}'

        # 生成會議記錄文件
        meeting_doc = Document()
        meeting_doc.add_heading(f'會議記錄: {meeting.title}', level=1)
        meeting_doc.add_paragraph(f'日期: {meeting.date}')
        meeting_doc.add_paragraph(f'時間: {meeting.start_time} - {meeting.end_time}')
        meeting_doc.add_paragraph(f'地點: {meeting.location}')
        meeting_doc.add_paragraph(f'參加人: {meeting.attendees}')
        meeting_doc.add_paragraph(f'記錄: {meeting.minutes or "無記錄"}')
        meeting_file = os.path.join(temp_dir, f'{base_name}_record.docx')
        meeting_doc.save(meeting_file)

        # 生成議程文件
        agenda_doc = Document()
        agenda_doc.add_heading(f'會議議程: {meeting.title}', level=1)
        for item in agenda_items:
            agenda_doc.add_paragraph(f'{item.item_number}. {item.item_title} - {item.description} '
                                  f'(負責人: {item.responsible_person}, 時間: {item.estimated_time})')
        agenda_file = os.path.join(temp_dir, f'{base_name}_agenda.docx')
        agenda_doc.save(agenda_file)

        # 生成附件文件（列出附件信息）
        attachment_doc = Document()
        attachment_doc.add_heading(f'會議附件: {meeting.title}', level=1)
        if attachments.exists():
            for attachment in attachments:
                attachment_doc.add_paragraph(f'文件名: {attachment.file.name} '
                                          f'(連結: {attachment.file.url})')
        else:
            attachment_doc.add_paragraph('無附件')
        attachment_file = os.path.join(temp_dir, f'{base_name}_attachments.docx')
        attachment_doc.save(attachment_file)

        # 壓縮為 ZIP 文件
        zip_path = default_storage.path(f'{base_name}_files.zip')
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(meeting_file, os.path.basename(meeting_file))
            zipf.write(agenda_file, os.path.basename(agenda_file))
            zipf.write(attachment_file, os.path.basename(attachment_file))

        # 回傳 ZIP 文件
        with open(zip_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="{base_name}_files.zip"'
            return response

    except ObjectDoesNotExist:
        return HttpResponse("會議不存在", status=404)
    except Exception as e:
        return HttpResponse(f"生成文件失敗: {str(e)}", status=500)
    finally:
        # 清理臨時文件和目錄
        if os.path.exists(temp_dir):
            for file in os.listdir(temp_dir):
                os.remove(os.path.join(temp_dir, file))
            os.rmdir(temp_dir)
        if os.path.exists(zip_path):
            os.remove(zip_path)